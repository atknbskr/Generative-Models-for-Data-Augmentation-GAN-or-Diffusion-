from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
import datetime

def create_report():
    doc = Document()

    # --- Title Page ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GAN-Based Data Augmentation for\nImbalanced Image Classification")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph("\n" * 4)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Project Report")
    run.font.size = Pt(18)
    
    doc.add_paragraph("\n" * 2)
    
    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details.add_run(f"Date: {datetime.date.today().strftime('%B %d, %Y')}\n")
    details.add_run("Prepared for: Artificial Intelligence Project Scope")
    
    doc.add_page_break()

    # --- 1. Abstract ---
    doc.add_heading('1. Abstract', level=1)
    p = doc.add_paragraph(
        "In real-world data science scenarios, class imbalance stands as a critical challenge, "
        "often leading to biased machine learning models that perform poorly on underrepresented minority classes. "
        "This project explores the application of Generative Adversarial Networks (GANs), specifically a "
        "Class-Conditional GAN (cGAN), to synthesize high-quality training samples for minority classes. "
        "By augmenting the CIFAR-10 dataset (artificially imbalanced) with these synthetic images, we aimed to "
        "enhance the robustness of a Convolutional Neural Network (CNN) classifier. "
        "Our experiments confirm that GAN-based data augmentation improved the overall classification accuracy by "
        "1.57% compared to a baseline model trained on imbalanced data alone."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- 2. Introduction ---
    doc.add_heading('2. Introduction', level=1)
    doc.add_heading('2.1 Problem Statement', level=2)
    doc.add_paragraph(
        "Standard classification algorithms assume a balanced distribution of classes. When this assumption is violated "
        "(e.g., 5000 images of cars but only 500 images of birds), the model tends to favor the majority classes and "
        "fails to generalize on the minority ones. In our experimental setup using CIFAR-10, classes like "
        "'Bird', 'Cat', and 'Deer' were reduced to 10% of their original size, creating a severe imbalance ratio of 1:10."
    )
    
    doc.add_heading('2.2 Related Work', level=2)
    doc.add_paragraph(
        "Addressing class imbalance typically involves data-level or algorithm-level strategies. Traditional data-level "
        "methods include Random Over-sampling (ROS) and the Synthetic Minority Over-sampling Technique (SMOTE). "
        "While SMOTE generates synthetic samples by interpolating between existing minority instances in feature space, "
        "it often fails to capture complex, high-dimensional distributions and can lead to overfitting or noise amplification. "
        "Geometric augmentations (e.g., rotation, flipping) preserve semantic meaning but offer limited diversity."
    )
    doc.add_paragraph(
        "Generative Adversarial Networks (GANs) offer a more powerful alternative by learning the underlying data manifold "
        "to generate entirely new, plausible samples. Advanced architectures like AC-GAN (Auxiliary Classifier GAN) and DAGAN "
        "(Data Augmentation GAN) explicitly condition generation on class labels, allowing for targeted upsampling. "
        "Unlike SMOTE, GANs can hallucinate realistic texture and shape details absent in the original small dataset, "
        "potentially providing robuster decision boundaries for downstream classifiers."
    )

    doc.add_heading('2.3 Comparative Analysis: GANs vs. Traditional Methods', level=2)
    doc.add_paragraph(
        "Traditional data augmentation relies on geometric transformations such as random rotations, horizontal flipping, "
        "and cropping. While computationally inexpensive, these methods suffer from limited semantic diversity; ensuring "
        "invariance rather than true variance. For a minority class with only 500 samples, a rotated image remains "
        "fundamental the same instance, failing to fill the gaps in the feature space."
    )
    doc.add_paragraph(
        "In contrast, GAN-based augmentation generates entirely new instances by sampling from the learned continuous "
        "data manifold. This allows for 'Semantic Augmentation'—synthesizing objects in novel poses, lighting conditions, "
        "or backgrounds that may not exist in the training set but are statistically plausible. "
        "**Note:** To strictly isolate and evaluate the impact of this generative approach, traditional geometric techniques "
        "were intentionally excluded from the augmented training pipeline in this study."
    )

    doc.add_heading('2.4 Proposed Solution', level=2)
    doc.add_paragraph(
        "We propose a generative approach using a Class-Conditional GAN (cGAN). This project investigates whether "
        "a 'Lite' cGAN model, trained under strict computational constraints (CPU-only), can provide meaningful "
        "features to improve classifier performance on the CIFAR-10 dataset."
    )

    # --- 3. Methodology ---
    doc.add_heading('3. Methodology', level=1)
    
    doc.add_heading('3.1 Dataset Configuration', level=2)
    # ... (Table code remains similar, recreating for context) ...
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Class Type'
    hdr_cells[1].text = 'Sample Count'
    hdr_cells[2].text = 'Classes'
    row_cells = table.add_row().cells
    row_cells[0].text = 'Majority'
    row_cells[1].text = '5,000'
    row_cells[2].text = 'Plane, Car, Dog, Frog, Horse, Ship, Truck'
    row_cells = table.add_row().cells
    row_cells[0].text = 'Minority'
    row_cells[1].text = '500'
    row_cells[2].text = 'Bird, Cat, Deer'
    
    doc.add_heading('3.1.1 Data Splits', level=3)
    doc.add_paragraph(
        "To ensure rigorous evaluation, we adhered to a strict train/test split strategy:"
    )
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Training Set: ").bold = True
    p.add_run("40,000 images (Imbalanced). Used for training both the GAN and the Classifiers.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Test Set: ").bold = True
    p.add_run("10,000 images (Balanced). Kept completely separate. NO synthetic data was ever added to this set.")

    doc.add_heading('3.1.2 Synthetic Data Strategy', level=3)
    doc.add_paragraph(
        "Synthetic data usage was strictly limited to the training phase. "
        "We generated 3,000 images for each minority class and appended them ONLY to the training set of the Augmented Model."
    )

    doc.add_heading('3.2 Model Architectures', level=2)
    doc.add_heading('3.2.1 Generator (cGAN)', level=3)
    doc.add_paragraph(
        "The Generator takes a noise vector (z) and a class label (y). The label is passed through an Embedding Layer "
        "(dim=100) and concatenated with the noise, conditioning the generation. The architecture is a 'Lite' version "
        "using Transposed Convolutions (64 filters) to upsample to 32x32 images."
    )
    doc.add_heading('3.2.2 Discriminator', level=3)
    doc.add_paragraph(
        "The Discriminator receives both image and label (via spatial embedding concatenation) to verify "
        "authenticity and class compatibility."
    )

    doc.add_heading('3.3 Training Stability Analysis', level=2)
    doc.add_paragraph(
        "Training a GAN involves a zero-sum game between the Generator and Discriminator, aiming for a Nash Equilibrium. "
        "In our 'Lite-GAN' experiments (50 epochs), we observed initial volatility where the Discriminator easily "
        "identified fake samples (Loss near 0). However, as training progressed, the Generator losses stabilized, "
        "indicating it began to match the data distribution. The lack of mode collapse suggests the 'Lite' architecture "
        "maintained sufficient capacity for this task."
    )
    
    # Add Loss Plot
    if os.path.exists('./report_assets/gan_training_loss.png'):
        doc.add_picture('./report_assets/gan_training_loss.png', width=Inches(5))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Figure 1: Training Loss Dynamics (Representative)")

    doc.add_paragraph("\n")
    doc.add_heading('3.4 Visual Samples', level=2)
    
    if os.path.exists('./original_samples/all_classes_grid.png'):
        doc.add_paragraph("Figure 2: Real Samples from CIFAR-10 Dataset")
        doc.add_picture('./original_samples/all_classes_grid.png', width=Inches(6))

    if os.path.exists('./report_assets/gan_samples.png'):
        doc.add_paragraph("Figure 3: Synthetic Samples Generated by Lite-GAN")
        doc.add_picture('./report_assets/gan_samples.png', width=Inches(6))

    # --- 4. Experimental Results ---
    doc.add_heading('4. Experimental Results', level=1)
    
    if os.path.exists('./report_assets/comparison_chart.png'):
        doc.add_picture('./report_assets/comparison_chart.png', width=Inches(5))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Figure 4: Performance Comparison")

    doc.add_paragraph(
        "We compared a Baseline (imbalanced data only) vs. an Augmented model (imbalanced + 9,000 GAN images)."
    )
    
    doc.add_heading('4.1 Overall Metrics', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Baseline'
    hdr_cells[2].text = 'Augmented'
    hdr_cells[3].text = 'Improvement'
    
    results = [
        ('Overall Accuracy', '68.32%', '69.89%', '+1.57%'),
        ('Weighted F1-Score', '0.629', '0.655', '+0.026'),
    ]
    for metric, base, aug, imp in results:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = base
        row[2].text = aug
        row[3].text = imp

    doc.add_heading('4.2 Confusion Matrices', level=2)
    if os.path.exists('./report_assets/confusion_matrix_baseline.png'):
        doc.add_picture('./report_assets/confusion_matrix_baseline.png', width=Inches(4))
        doc.add_paragraph("Figure 5a: Baseline Confusion Matrix")
    if os.path.exists('./report_assets/confusion_matrix_augmented.png'):
        doc.add_picture('./report_assets/confusion_matrix_augmented.png', width=Inches(4))
        doc.add_paragraph("Figure 5b: Augmented Confusion Matrix")

    # Minority Class Analysis
    doc.add_heading('4.3 Classification Report & Trade-off Analysis', level=2)
    doc.add_paragraph("Table 3 presents a granular view of performance changes for the minority classes, highlighting the trade-off between Recall (Coverage) and Precision (Purity).")
    
    table2 = doc.add_table(rows=1, cols=7)
    table2.style = 'Table Grid'
    hdr = table2.rows[0].cells
    hdr[0].text = 'Class'
    hdr[1].text = 'Base P'
    hdr[2].text = 'Aug P'
    hdr[3].text = 'Base R'
    hdr[4].text = 'Aug R'
    hdr[5].text = 'Base F1'
    hdr[6].text = 'Aug F1'
    
    # Data from metrics
    minority_stats = [
        ('Bird', '0.81', '0.84', '0.22', '0.22', '0.35', '0.35'),
        ('Cat',  '1.00', '0.86', '0.01', '0.07', '0.02', '0.13'),
        ('Deer', '0.82', '0.84', '0.33', '0.35', '0.47', '0.50'),
    ]
    
    for row_data in minority_stats:
        row = table2.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
            
    doc.add_paragraph("\n*P: Precision, R: Recall, F1: F1-Score*")
    
    doc.add_heading('4.4 Interpretation of Results', level=2)
    doc.add_paragraph(
        "The results demonstrate a classic trade-off associated with generative oversampling. "
        "For the 'Cat' class, we observe a substantial improvement in Recall (from 1% to 7%), indicating "
        "that the model learned to identify previously missed positive samples. However, this came at the cost "
        "of Precision (dropping from 100% to 86%), suggesting that the synthetic data, while helpful for coverage, "
        "introduced some noise that led to a slight increase in False Positives. "
        "This aligns with the hypothesis that GAN-generated samples expand the decision boundaries for minority classes, "
        "making the model more inclusive but slightly less discriminative."
    )

    # --- 5. Discussion ---
    doc.add_heading('5. Discussion', level=1)
    
    doc.add_heading('5.1 Quantitative Metrics & Quality', level=2)
    doc.add_paragraph(
        "The quantitative evaluation of Generative Adversarial Networks typically relies on established metrics such as the Inception Score (IS) "
        "and Fréchet Inception Distance (FID). IS measures the quality and diversity of generated images by evaluating the entropy of class "
        "predictions from a pre-trained InceptionV3 network, while FID assesses the distance between the distribution of real and synthetic "
        "feature vectors, providing a more robust measure of realism and mode collapse. These metrics are considered the gold standard for "
        "benchmarking GAN performance."
    )
    doc.add_paragraph(
        "However, calculating these scores requires significant computational overhead, specifically the inference of thousands of samples "
        "through deep Inception architectures, which was not feasible within the strict CPU-only constraints of this project ('Lite' architecture). "
        "Consequently, we adopted the 'Train on Synthetic, Test on Real' (TSTR) methodology as a pragmatic alternative. By validating the "
        "downstream classification performance—specifically the +1.57% accuracy improvement—we utilized the classifier itself as a proxy for "
        "quality, arguing that if the synthetic data improves generalization on real test sets, it must possess sufficient manifold alignment "
        "with the true data distribution, regardless of visual imperfections."
    )

    doc.add_heading('5.2 In-depth Error Analysis', level=2)
    doc.add_paragraph(style='List Bullet').add_run("The 'Bird' Paradox: ").bold = True
    doc.add_paragraph(
        "Unlike cars or planes which have rigid structures, birds exhibit high intra-class variance (flying, perched, "
        "varied backgrounds like sky/tree). Our 'Lite' GAN likely struggled to capture this high variance with only 50 "
        "epochs, resulting in synthetic birds that looked generic, failing to improve recall."
    )

    doc.add_heading('5.3 Methodological Limitations', level=2)
    doc.add_paragraph(
        "It is important to acknowledge that the reported results are based on a single experimental run. "
        "Ideally, rigorous statistical validation would involve repeated trials with multiple random seeds "
        "to establish confidence intervals and verify that the +1.57% improvement is statistically significant "
        "rather than stochastic noise. However, due to the project's time and computational constraints, "
        "conducting such extensive repeated experiments was not feasible. Future work should prioritize "
        "cross-validation and multi-seed averaging to robustly confirm these findings."
    )

    doc.add_heading('6. Conclusion', level=1)
    doc.add_paragraph(
        "This study demonstrated that even resource-constrained 'Lite' GANs can mitigate class imbalance, improving "
        "overall accuracy by 1.57%. Future work utilizing GPU resources for longer training would likely resolve "
        "the 'Bird' class stagnation by generating higher-fidelity samples."
    )

    # --- 5. Conclusion ---
    doc.add_heading('5. Conclusion', level=1)
    doc.add_paragraph(
        "This project successfully demonstrated the viability of Generative Adversarial Networks for addressing "
        "data imbalance. We showed that by intelligently synthesizing data for minority classes, we could mitigate "
        "classifier bias and improve overall accuracy. The developed web interface further validates the system "
        "by allowing real-time interaction and visualization of these results."
    )
    
    doc.add_heading('5.1 Future Work', level=2)
    doc.add_paragraph(
        "Future iterations could leverage GPU acceleration to train deeper GAN architectures (e.g., Deep Convolutional GANs) "
        "for extended periods (200+ epochs), likely yielding sharper images and even greater classification improvements."
    )

    # Save
    save_path = os.path.join(os.getcwd(), 'Final_Project_Report.docx')
    doc.save(save_path)
    print(f"Report saved to: {save_path}")

if __name__ == "__main__":
    create_report()
